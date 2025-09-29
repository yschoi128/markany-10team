#!/usr/bin/env python3
"""
AI 식단 코치 프로젝트 문서 생성기
프로젝트 분석 후 6개 문서를 docx 형식으로 생성
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

class DocumentGenerator:
    def __init__(self, project_path):
        self.project_path = Path(project_path)
        self.output_dir = Path("generated_docs")
        self.output_dir.mkdir(exist_ok=True)
        
    def create_project_overview(self):
        """프로젝트 개요서 생성"""
        doc = Document()
        
        # 제목
        title = doc.add_heading('AI 식단 코치 프로젝트 개요서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 기본 정보
        doc.add_heading('1. 프로젝트 기본 정보', level=1)
        
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        info_data = [
            ('프로젝트명', 'AI 식단 코치 (AI Diet Coach)'),
            ('개발팀', 'Markany 10팀'),
            ('개발기간', '2024년 해커톤'),
            ('기술스택', 'Python, FastAPI, AWS (Bedrock, DynamoDB, S3)'),
            ('프로젝트 유형', 'AI 기반 개인 맞춤형 식단 관리 솔루션'),
            ('배포환경', 'AWS 클라우드')
        ]
        
        for i, (key, value) in enumerate(info_data):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value
        
        # 프로젝트 개요
        doc.add_heading('2. 프로젝트 개요', level=1)
        doc.add_paragraph(
            'AWS 기반의 개인 맞춤형 식단 관리 AI 솔루션입니다. '
            '회식이나 일상 식사 사진을 촬영하면 음식 종류와 섭취량을 자동 분석하고, '
            '개인의 건강 목표와 운동 선호도를 바탕으로 맞춤형 식단과 운동을 처방합니다. '
            '음성 기반 대화형 인터페이스를 통해 AI PT가 일일/주간 식습관을 분석하고 '
            '실시간 조언을 제공하며, 스케줄러 연동으로 회식 등 예정된 식사까지 관리하여 '
            '건강한 라이프스타일을 지원합니다.'
        )
        
        # 주요 기능
        doc.add_heading('3. 주요 기능', level=1)
        
        features = [
            '식사 이미지 분석: 사진 촬영 시 음식 종류, 칼로리, 영양소 자동 분석',
            'AI PT 코칭: 개인 목표 기반 맞춤형 식단/운동 처방 및 실시간 조언',
            '스케줄 연동 관리: 회식 등 예정된 식사 스케줄링 및 사전/사후 관리',
            '대화형 인터페이스: 음성/텍스트 기반 자연어 대화',
            '개인화 추천: 사용자 프로필 기반 맞춤형 식단 및 운동 추천'
        ]
        
        for feature in features:
            p = doc.add_paragraph(feature, style='List Bullet')
        
        # 기술적 특징
        doc.add_heading('4. 기술적 특징', level=1)
        
        tech_features = [
            'AWS Bedrock을 활용한 멀티모달 AI 분석',
            'FastAPI 기반 RESTful API 설계',
            'DynamoDB를 활용한 NoSQL 데이터 관리',
            'S3를 활용한 이미지 저장 및 관리',
            'Pydantic을 활용한 타입 안전성 보장',
            '비동기 처리를 통한 성능 최적화'
        ]
        
        for tech in tech_features:
            p = doc.add_paragraph(tech, style='List Bullet')
        
        # 기대 효과
        doc.add_heading('5. 기대 효과', level=1)
        doc.add_paragraph(
            '• 개인 맞춤형 식단 관리를 통한 건강 개선\n'
            '• AI 기반 자동 분석으로 사용자 편의성 향상\n'
            '• 실시간 코칭을 통한 지속적인 동기 부여\n'
            '• 스케줄 연동을 통한 체계적인 식습관 관리\n'
            '• 데이터 기반 개인화 추천으로 효과적인 건강 관리'
        )
        
        doc.save(self.output_dir / '01_프로젝트_개요서.docx')
        print("✅ 프로젝트 개요서 생성 완료")

    def create_requirements_document(self):
        """요구사항 정의서 생성"""
        doc = Document()
        
        # 제목
        title = doc.add_heading('AI 식단 코치 요구사항 정의서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. 기능 요구사항
        doc.add_heading('1. 기능 요구사항', level=1)
        
        # 1.1 사용자 관리
        doc.add_heading('1.1 사용자 관리', level=2)
        
        user_reqs = [
            'REQ-001: 사용자는 개인 프로필(나이, 성별, 신장, 체중, 건강목표)을 등록할 수 있다',
            'REQ-002: 사용자는 선호/비선호 운동 종류를 설정할 수 있다',
            'REQ-003: 사용자는 식이 제한사항(알레르기, 종교적 제약 등)을 설정할 수 있다',
            'REQ-004: 시스템은 사용자 정보를 기반으로 목표 칼로리를 자동 계산한다'
        ]
        
        for req in user_reqs:
            doc.add_paragraph(req, style='List Bullet')
        
        # 1.2 식사 이미지 분석
        doc.add_heading('1.2 식사 이미지 분석', level=2)
        
        image_reqs = [
            'REQ-005: 사용자는 식사 사진을 업로드할 수 있다',
            'REQ-006: 시스템은 이미지에서 음식 종류를 자동 인식한다',
            'REQ-007: 시스템은 음식별 예상 칼로리와 영양소를 계산한다',
            'REQ-008: 시스템은 함께 식사한 인원 수를 고려하여 1인분 섭취량을 계산한다',
            'REQ-009: 시스템은 분석 결과의 신뢰도를 제공한다'
        ]
        
        for req in image_reqs:
            doc.add_paragraph(req, style='List Bullet')
        
        # 1.3 AI 코칭
        doc.add_heading('1.3 AI 코칭', level=2)
        
        coaching_reqs = [
            'REQ-010: 시스템은 일일 식단 분석 결과를 제공한다',
            'REQ-011: 시스템은 개인 목표 대비 진행상황을 분석한다',
            'REQ-012: 시스템은 맞춤형 식단 추천을 제공한다',
            'REQ-013: 시스템은 개인 선호도 기반 운동 추천을 제공한다',
            'REQ-014: 시스템은 음성/텍스트 기반 대화형 인터페이스를 제공한다'
        ]
        
        for req in coaching_reqs:
            doc.add_paragraph(req, style='List Bullet')
        
        # 2. 비기능 요구사항
        doc.add_heading('2. 비기능 요구사항', level=1)
        
        # 2.1 성능 요구사항
        doc.add_heading('2.1 성능 요구사항', level=2)
        
        performance_reqs = [
            'NFR-001: 이미지 분석 응답시간은 10초 이내여야 한다',
            'NFR-002: API 응답시간은 3초 이내여야 한다',
            'NFR-003: 시스템은 동시 사용자 100명을 지원해야 한다'
        ]
        
        for req in performance_reqs:
            doc.add_paragraph(req, style='List Bullet')
        
        doc.save(self.output_dir / '02_요구사항_정의서.docx')
        print("✅ 요구사항 정의서 생성 완료")

    def create_architecture_document(self):
        """시스템 아키텍처 다이어그램 생성"""
        doc = Document()
        
        # 제목
        title = doc.add_heading('AI 식단 코치 시스템 아키텍처', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. 전체 아키텍처 개요
        doc.add_heading('1. 전체 아키텍처 개요', level=1)
        doc.add_paragraph(
            'AI 식단 코치는 마이크로서비스 아키텍처를 기반으로 하며, '
            'AWS 클라우드 서비스를 활용한 서버리스 및 관리형 서비스로 구성됩니다.'
        )
        
        # 2. 아키텍처 구성요소
        doc.add_heading('2. 아키텍처 구성요소', level=1)
        
        # 2.1 프레젠테이션 계층
        doc.add_heading('2.1 프레젠테이션 계층', level=2)
        presentation_components = [
            'FastAPI 웹 애플리케이션 서버',
            'RESTful API 엔드포인트',
            'Swagger/OpenAPI 문서화',
            'CORS 미들웨어'
        ]
        for comp in presentation_components:
            doc.add_paragraph(comp, style='List Bullet')
        
        # 2.2 비즈니스 로직 계층
        doc.add_heading('2.2 비즈니스 로직 계층', level=2)
        business_components = [
            'Food Analysis Pipeline: 식사 이미지 분석 파이프라인',
            'Coaching Pipeline: AI 코칭 및 추천 파이프라인',
            'User Management Service: 사용자 관리 서비스',
            'Schedule Management Service: 스케줄 관리 서비스'
        ]
        for comp in business_components:
            doc.add_paragraph(comp, style='List Bullet')
        
        # 2.3 데이터 계층
        doc.add_heading('2.3 데이터 계층', level=2)
        data_components = [
            'Amazon DynamoDB: NoSQL 데이터베이스',
            'Amazon S3: 이미지 및 파일 저장소',
            'AWS Bedrock: AI/ML 모델 서비스'
        ]
        for comp in data_components:
            doc.add_paragraph(comp, style='List Bullet')
        
        # 3. 데이터 플로우
        doc.add_heading('3. 데이터 플로우', level=1)
        
        doc.add_heading('3.1 식사 이미지 분석 플로우', level=2)
        flow_steps = [
            '1. 사용자가 식사 이미지를 업로드',
            '2. 이미지가 S3 버킷에 저장',
            '3. Bedrock 이미지 분석 모델로 음식 인식',
            '4. 영양소 및 칼로리 계산',
            '5. 분석 결과를 DynamoDB에 저장',
            '6. 사용자에게 분석 결과 반환'
        ]
        for step in flow_steps:
            doc.add_paragraph(step, style='List Number')
        
        # 4. 기술 스택
        doc.add_heading('4. 기술 스택', level=1)
        
        tech_table = doc.add_table(rows=6, cols=2)
        tech_table.style = 'Table Grid'
        
        tech_data = [
            ('백엔드 프레임워크', 'FastAPI (Python)'),
            ('AI/ML 서비스', 'AWS Bedrock (Claude, Titan)'),
            ('데이터베이스', 'Amazon DynamoDB'),
            ('파일 저장소', 'Amazon S3'),
            ('배포 환경', 'AWS Lambda / EC2'),
            ('모니터링', 'CloudWatch')
        ]
        
        for i, (category, tech) in enumerate(tech_data):
            tech_table.cell(i, 0).text = category
            tech_table.cell(i, 1).text = tech
        
        # 5. 보안 아키텍처
        doc.add_heading('5. 보안 아키텍처', level=1)
        security_features = [
            'HTTPS 통신 강제',
            'AWS IAM 역할 기반 접근 제어',
            'DynamoDB 암호화 저장',
            'S3 버킷 정책 및 암호화',
            'API 키 기반 인증'
        ]
        for feature in security_features:
            doc.add_paragraph(feature, style='List Bullet')
        
        doc.save(self.output_dir / '03_시스템_아키텍처.docx')
        print("✅ 시스템 아키텍처 문서 생성 완료")

    def create_database_design(self):
        """DB 설계서 생성"""
        doc = Document()
        
        # 제목
        title = doc.add_heading('AI 식단 코치 데이터베이스 설계서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. 데이터베이스 개요
        doc.add_heading('1. 데이터베이스 개요', level=1)
        doc.add_paragraph(
            'AI 식단 코치는 Amazon DynamoDB를 주 데이터베이스로 사용하며, '
            'NoSQL 기반의 확장 가능한 데이터 모델을 채택합니다. '
            '이미지 파일은 Amazon S3에 저장되며, DynamoDB에는 메타데이터만 저장됩니다.'
        )
        
        # 2. 테이블 설계
        doc.add_heading('2. 테이블 설계', level=1)
        
        # 2.1 사용자 프로필 테이블
        doc.add_heading('2.1 user_profiles 테이블', level=2)
        doc.add_paragraph('사용자의 기본 정보와 건강 목표를 저장하는 테이블')
        
        user_table = doc.add_table(rows=13, cols=4)
        user_table.style = 'Table Grid'
        
        # 헤더
        user_table.cell(0, 0).text = '필드명'
        user_table.cell(0, 1).text = '타입'
        user_table.cell(0, 2).text = '설명'
        user_table.cell(0, 3).text = '제약조건'
        
        user_fields = [
            ('user_id', 'String', '사용자 고유 ID', 'Partition Key'),
            ('name', 'String', '사용자 이름', 'Required'),
            ('age', 'Number', '나이', 'Required'),
            ('gender', 'String', '성별', 'Required'),
            ('height', 'Number', '신장(cm)', 'Required'),
            ('weight', 'Number', '체중(kg)', 'Required'),
            ('health_goal', 'String', '건강 목표', 'Required'),
            ('preferred_exercises', 'List', '선호 운동', 'Optional'),
            ('activity_level', 'String', '활동량', 'Required'),
            ('target_calories', 'Number', '목표 칼로리', 'Optional'),
            ('created_at', 'String', '생성일시', 'Required'),
            ('updated_at', 'String', '수정일시', 'Required')
        ]
        
        for i, (field, type_name, desc, constraint) in enumerate(user_fields, 1):
            user_table.cell(i, 0).text = field
            user_table.cell(i, 1).text = type_name
            user_table.cell(i, 2).text = desc
            user_table.cell(i, 3).text = constraint
        
        # 2.2 식사 기록 테이블
        doc.add_heading('2.2 diet_records 테이블', level=2)
        doc.add_paragraph('사용자의 식사 기록과 영양소 정보를 저장하는 테이블')
        
        diet_table = doc.add_table(rows=12, cols=4)
        diet_table.style = 'Table Grid'
        
        # 헤더
        diet_table.cell(0, 0).text = '필드명'
        diet_table.cell(0, 1).text = '타입'
        diet_table.cell(0, 2).text = '설명'
        diet_table.cell(0, 3).text = '제약조건'
        
        diet_fields = [
            ('user_id', 'String', '사용자 ID', 'Partition Key'),
            ('meal_id', 'String', '식사 ID', 'Sort Key'),
            ('timestamp', 'String', '식사 시간', 'Required'),
            ('meal_type', 'String', '식사 종류', 'Required'),
            ('image_url', 'String', 'S3 이미지 URL', 'Optional'),
            ('foods', 'List', '음식 목록', 'Required'),
            ('total_calories', 'Number', '총 칼로리', 'Required'),
            ('total_carbs', 'Number', '총 탄수화물(g)', 'Required'),
            ('total_protein', 'Number', '총 단백질(g)', 'Required'),
            ('total_fat', 'Number', '총 지방(g)', 'Required'),
            ('notes', 'String', '메모', 'Optional')
        ]
        
        for i, (field, type_name, desc, constraint) in enumerate(diet_fields, 1):
            diet_table.cell(i, 0).text = field
            diet_table.cell(i, 1).text = type_name
            diet_table.cell(i, 2).text = desc
            diet_table.cell(i, 3).text = constraint
        
        # 2.3 스케줄 테이블
        doc.add_heading('2.3 schedule_records 테이블', level=2)
        doc.add_paragraph('사용자의 식사 일정을 저장하는 테이블')
        
        schedule_table = doc.add_table(rows=9, cols=4)
        schedule_table.style = 'Table Grid'
        
        # 헤더
        schedule_table.cell(0, 0).text = '필드명'
        schedule_table.cell(0, 1).text = '타입'
        schedule_table.cell(0, 2).text = '설명'
        schedule_table.cell(0, 3).text = '제약조건'
        
        schedule_fields = [
            ('event_id', 'String', '이벤트 ID', 'Partition Key'),
            ('user_id', 'String', '사용자 ID', 'GSI Partition Key'),
            ('title', 'String', '이벤트 제목', 'Required'),
            ('event_type', 'String', '이벤트 종류', 'Required'),
            ('start_time', 'String', '시작 시간', 'Required'),
            ('location', 'String', '장소', 'Optional'),
            ('participants', 'Number', '참석 인원', 'Optional'),
            ('is_processed', 'Boolean', '처리 완료 여부', 'Required')
        ]
        
        for i, (field, type_name, desc, constraint) in enumerate(schedule_fields, 1):
            schedule_table.cell(i, 0).text = field
            schedule_table.cell(i, 1).text = type_name
            schedule_table.cell(i, 2).text = desc
            schedule_table.cell(i, 3).text = constraint
        
        # 3. 인덱스 설계
        doc.add_heading('3. 인덱스 설계', level=1)
        
        doc.add_heading('3.1 Global Secondary Index (GSI)', level=2)
        gsi_info = [
            'user_id-timestamp-index: 사용자별 시간순 식사 기록 조회',
            'user_id-meal_type-index: 사용자별 식사 종류별 조회',
            'user_id-date-index: 사용자별 일별 식사 기록 조회'
        ]
        for info in gsi_info:
            doc.add_paragraph(info, style='List Bullet')
        
        # 4. 데이터 관계도 (ERD 설명)
        doc.add_heading('4. 데이터 관계도', level=1)
        doc.add_paragraph(
            'DynamoDB는 NoSQL 데이터베이스이므로 전통적인 관계형 모델과는 다르지만, '
            '논리적 관계는 다음과 같습니다:'
        )
        
        relationships = [
            '사용자(user_profiles) 1 : N 식사기록(diet_records)',
            '사용자(user_profiles) 1 : N 스케줄(schedule_records)',
            '식사기록(diet_records) 1 : 1 이미지파일(S3)'
        ]
        for rel in relationships:
            doc.add_paragraph(rel, style='List Bullet')
        
        doc.save(self.output_dir / '04_DB_설계서.docx')
        print("✅ DB 설계서 생성 완료")

    def create_api_specification(self):
        """API 명세서 생성"""
        doc = Document()
        
        # 제목
        title = doc.add_heading('AI 식단 코치 API 명세서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. API 개요
        doc.add_heading('1. API 개요', level=1)
        doc.add_paragraph(
            'AI 식단 코치는 RESTful API를 제공하며, JSON 형식으로 데이터를 주고받습니다. '
            '모든 API는 HTTPS 통신을 사용하며, 적절한 HTTP 상태 코드를 반환합니다.'
        )
        
        # 기본 정보 테이블
        api_info_table = doc.add_table(rows=4, cols=2)
        api_info_table.style = 'Table Grid'
        
        api_info = [
            ('Base URL', 'https://api.ai-diet-coach.com'),
            ('Protocol', 'HTTPS'),
            ('Data Format', 'JSON'),
            ('Authentication', 'API Key')
        ]
        
        for i, (key, value) in enumerate(api_info):
            api_info_table.cell(i, 0).text = key
            api_info_table.cell(i, 1).text = value
        
        # 2. 사용자 관리 API
        doc.add_heading('2. 사용자 관리 API', level=1)
        
        # 2.1 사용자 프로필 생성
        doc.add_heading('2.1 사용자 프로필 생성', level=2)
        doc.add_paragraph('POST /users')
        doc.add_paragraph('새로운 사용자 프로필을 생성합니다.')
        
        # 요청 예시
        doc.add_heading('요청 예시:', level=3)
        request_example = '''
{
  "name": "홍길동",
  "age": 30,
  "gender": "male",
  "height": 175,
  "weight": 70,
  "health_goal": "weight_loss",
  "preferred_exercises": ["gym", "running"],
  "activity_level": "moderate"
}
        '''
        p = doc.add_paragraph()
        p.add_run(request_example).font.name = 'Courier New'
        
        # 응답 예시
        doc.add_heading('응답 예시:', level=3)
        response_example = '''
{
  "success": true,
  "message": "사용자 프로필이 성공적으로 생성되었습니다.",
  "data": {
    "user_id": "user_12345"
  }
}
        '''
        p = doc.add_paragraph()
        p.add_run(response_example).font.name = 'Courier New'
        
        # 2.2 사용자 프로필 조회
        doc.add_heading('2.2 사용자 프로필 조회', level=2)
        doc.add_paragraph('GET /users/{user_id}')
        doc.add_paragraph('특정 사용자의 프로필 정보를 조회합니다.')
        
        # 3. 식사 분석 API
        doc.add_heading('3. 식사 분석 API', level=1)
        
        # 3.1 식사 이미지 분석
        doc.add_heading('3.1 식사 이미지 분석', level=2)
        doc.add_paragraph('POST /meals/analyze')
        doc.add_paragraph('업로드된 식사 이미지를 분석하여 영양소 정보를 제공합니다.')
        
        # 요청 파라미터
        doc.add_heading('요청 파라미터 (Form Data):', level=3)
        meal_params = [
            'user_id (string): 사용자 ID',
            'meal_type (string): 식사 종류 (아침/점심/저녁/간식)',
            'people_count (integer): 함께 식사한 인원 수',
            'image (file): 식사 이미지 파일',
            'notes (string, optional): 추가 메모'
        ]
        for param in meal_params:
            doc.add_paragraph(param, style='List Bullet')
        
        # 3.2 식사 기록 조회
        doc.add_heading('3.2 식사 기록 조회', level=2)
        doc.add_paragraph('GET /meals/{user_id}?days=7&limit=50')
        doc.add_paragraph('사용자의 식사 기록을 조회합니다.')
        
        # 4. 코칭 API
        doc.add_heading('4. 코칭 API', level=1)
        
        # 4.1 일일 코칭 메시지 생성
        doc.add_heading('4.1 일일 코칭 메시지 생성', level=2)
        doc.add_paragraph('POST /coaching/daily/{user_id}')
        doc.add_paragraph('사용자의 일일 식단을 분석하여 맞춤형 코칭 메시지를 생성합니다.')
        
        # 4.2 AI 코치와 대화
        doc.add_heading('4.2 AI 코치와 대화', level=2)
        doc.add_paragraph('POST /coaching/chat/{user_id}')
        doc.add_paragraph('사용자와 AI 코치 간의 대화형 인터페이스를 제공합니다.')
        
        # 5. 리포트 API
        doc.add_heading('5. 리포트 API', level=1)
        
        # 5.1 주간 리포트
        doc.add_heading('5.1 주간 리포트', level=2)
        doc.add_paragraph('GET /reports/weekly/{user_id}')
        doc.add_paragraph('사용자의 주간 식단 및 건강 리포트를 생성합니다.')
        
        # 5.2 일일 영양소 현황
        doc.add_heading('5.2 일일 영양소 현황', level=2)
        doc.add_paragraph('GET /nutrition/daily/{user_id}?date=2024-01-01')
        doc.add_paragraph('특정 날짜의 영양소 섭취 현황을 조회합니다.')
        
        # 6. 에러 코드
        doc.add_heading('6. 에러 코드', level=1)
        
        error_table = doc.add_table(rows=6, cols=2)
        error_table.style = 'Table Grid'
        
        error_table.cell(0, 0).text = 'HTTP 상태 코드'
        error_table.cell(0, 1).text = '설명'
        
        error_codes = [
            ('200', '성공'),
            ('400', '잘못된 요청'),
            ('401', '인증 실패'),
            ('404', '리소스를 찾을 수 없음'),
            ('500', '서버 내부 오류')
        ]
        
        for i, (code, desc) in enumerate(error_codes, 1):
            error_table.cell(i, 0).text = code
            error_table.cell(i, 1).text = desc
        
        doc.save(self.output_dir / '05_API_명세서.docx')
        print("✅ API 명세서 생성 완료")

    def create_readme_document(self):
        """README 문서 생성"""
        doc = Document()
        
        # 제목
        title = doc.add_heading('AI 식단 코치 README', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. 프로젝트 소개
        doc.add_heading('1. 프로젝트 소개', level=1)
        doc.add_paragraph(
            'AI 식단 코치는 AWS 기반의 개인 맞춤형 식단 관리 AI 솔루션입니다. '
            '식사 사진을 촬영하면 자동으로 음식을 인식하고 영양소를 분석하여, '
            '개인의 건강 목표에 맞는 맞춤형 코칭을 제공합니다.'
        )
        
        # 2. 주요 기능
        doc.add_heading('2. 주요 기능', level=1)
        features = [
            '🍽️ 식사 이미지 자동 분석 및 영양소 계산',
            '🤖 AI 기반 개인 맞춤형 식단 코칭',
            '📅 스케줄 연동 식사 관리',
            '💬 대화형 AI 코치 인터페이스',
            '📊 일일/주간 영양 리포트 생성'
        ]
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')
        
        # 3. 기술 스택
        doc.add_heading('3. 기술 스택', level=1)
        
        tech_table = doc.add_table(rows=7, cols=2)
        tech_table.style = 'Table Grid'
        
        tech_table.cell(0, 0).text = '분야'
        tech_table.cell(0, 1).text = '기술'
        
        tech_stack = [
            ('Backend', 'Python 3.12, FastAPI'),
            ('AI/ML', 'AWS Bedrock (Claude, Titan)'),
            ('Database', 'Amazon DynamoDB'),
            ('Storage', 'Amazon S3'),
            ('Deployment', 'AWS Lambda, EC2'),
            ('Monitoring', 'CloudWatch')
        ]
        
        for i, (category, tech) in enumerate(tech_stack, 1):
            tech_table.cell(i, 0).text = category
            tech_table.cell(i, 1).text = tech
        
        # 4. 시스템 요구사항
        doc.add_heading('4. 시스템 요구사항', level=1)
        requirements = [
            'Python 3.12 이상',
            'AWS 계정 및 자격 증명',
            'pip (Python 패키지 관리자)',
            '최소 2GB RAM',
            '인터넷 연결'
        ]
        for req in requirements:
            doc.add_paragraph(req, style='List Bullet')
        
        # 5. 설치 및 설정
        doc.add_heading('5. 설치 및 설정', level=1)
        
        doc.add_heading('5.1 프로젝트 클론', level=2)
        p = doc.add_paragraph()
        p.add_run('git clone https://github.com/yschoi128/markany-10team.git\ncd markany-10team').font.name = 'Courier New'
        
        doc.add_heading('5.2 가상환경 설정', level=2)
        venv_commands = '''
# 가상환경 생성
python -m venv venv

# 가상환경 활성화 (Linux/Mac)
source venv/bin/activate

# 가상환경 활성화 (Windows)
venv\\Scripts\\activate
        '''
        p = doc.add_paragraph()
        p.add_run(venv_commands).font.name = 'Courier New'
        
        doc.add_heading('5.3 의존성 설치', level=2)
        p = doc.add_paragraph()
        p.add_run('pip install -r requirements.txt').font.name = 'Courier New'
        
        doc.add_heading('5.4 환경 변수 설정', level=2)
        p = doc.add_paragraph()
        p.add_run('cp .env.example .env').font.name = 'Courier New'
        doc.add_paragraph('.env 파일을 편집하여 AWS 자격 증명을 설정하세요:')
        
        env_example = '''
AWS_ACCESS_KEY_ID=your_access_key
AWS_SECRET_ACCESS_KEY=your_secret_key
AWS_REGION=ap-northeast-2
S3_BUCKET_NAME=ai-diet-coach-images
DYNAMODB_DIET_TABLE=diet_records
BEDROCK_MODEL_ID=anthropic.claude-3-sonnet-20240229-v1:0
        '''
        p = doc.add_paragraph()
        p.add_run(env_example).font.name = 'Courier New'
        
        # 6. 실행 방법
        doc.add_heading('6. 실행 방법', level=1)
        
        doc.add_heading('6.1 개발 서버 실행', level=2)
        run_commands = '''
# 방법 1: 직접 실행
python run_agent.py

# 방법 2: uvicorn 사용
uvicorn src.main:app --reload --host 0.0.0.0 --port 8000
        '''
        p = doc.add_paragraph()
        p.add_run(run_commands).font.name = 'Courier New'
        
        doc.add_heading('6.2 API 테스트', level=2)
        doc.add_paragraph('서버 실행 후 다음 URL에서 API 문서를 확인할 수 있습니다:')
        api_urls = [
            'API 문서: http://localhost:8000/docs',
            'ReDoc: http://localhost:8000/redoc',
            '헬스체크: http://localhost:8000/health'
        ]
        for url in api_urls:
            doc.add_paragraph(url, style='List Bullet')
        
        # 7. 사용 예시
        doc.add_heading('7. 사용 예시', level=1)
        
        doc.add_heading('7.1 사용자 프로필 생성', level=2)
        user_example = '''
curl -X POST "http://localhost:8000/users" \\
  -H "Content-Type: application/json" \\
  -d '{
    "name": "홍길동",
    "age": 30,
    "gender": "male",
    "height": 175,
    "weight": 70,
    "health_goal": "weight_loss"
  }'
        '''
        p = doc.add_paragraph()
        p.add_run(user_example).font.name = 'Courier New'
        
        doc.add_heading('7.2 식사 이미지 분석', level=2)
        meal_example = '''
curl -X POST "http://localhost:8000/meals/analyze" \\
  -F "user_id=user123" \\
  -F "meal_type=점심" \\
  -F "people_count=1" \\
  -F "image=@meal_photo.jpg"
        '''
        p = doc.add_paragraph()
        p.add_run(meal_example).font.name = 'Courier New'
        
        # 8. 테스트
        doc.add_heading('8. 테스트', level=1)
        test_commands = '''
# 단위 테스트 실행
python -m pytest tests/

# 커버리지 포함 테스트
python -m pytest tests/ --cov=src

# 특정 테스트 실행
python -m pytest tests/test_food_analysis.py
        '''
        p = doc.add_paragraph()
        p.add_run(test_commands).font.name = 'Courier New'
        
        # 9. 배포
        doc.add_heading('9. 배포', level=1)
        doc.add_paragraph('AWS Lambda 또는 EC2를 통해 배포할 수 있습니다. 자세한 배포 가이드는 별도 문서를 참조하세요.')
        
        # 10. 문제 해결
        doc.add_heading('10. 문제 해결', level=1)
        troubleshooting = [
            'AWS 자격 증명 오류: .env 파일의 AWS 키 확인',
            'DynamoDB 테이블 없음: AWS 콘솔에서 테이블 생성 확인',
            'S3 버킷 접근 오류: 버킷 정책 및 권한 확인',
            'Bedrock 모델 오류: 모델 ID 및 리전 확인'
        ]
        for item in troubleshooting:
            doc.add_paragraph(item, style='List Bullet')
        
        # 11. 기여하기
        doc.add_heading('11. 기여하기', level=1)
        doc.add_paragraph(
            '1. Fork the repository\n'
            '2. Create a feature branch\n'
            '3. Commit your changes\n'
            '4. Push to the branch\n'
            '5. Open a Pull Request'
        )
        
        doc.save(self.output_dir / '06_README.docx')
        print("✅ README 문서 생성 완료")

    def generate_all_documents(self):
        """모든 문서 생성"""
        print("🚀 AI 식단 코치 문서 생성을 시작합니다...")
        
        self.create_project_overview()
        self.create_requirements_document()
        self.create_architecture_document()
        self.create_database_design()
        self.create_api_specification()
        self.create_readme_document()
        
        print(f"\n✅ 모든 문서가 '{self.output_dir}' 폴더에 생성되었습니다!")
        print("생성된 문서:")
        print("  📋 01_프로젝트_개요서.docx")
        print("  📝 02_요구사항_정의서.docx")
        print("  🏗️ 03_시스템_아키텍처.docx")
        print("  🗄️ 04_DB_설계서.docx")
        print("  🔌 05_API_명세서.docx")
        print("  📖 06_README.docx")

if __name__ == "__main__":
    project_path = "/home/sunhk/q/markany-10team"
    generator = DocumentGenerator(project_path)
    generator.generate_all_documents()
